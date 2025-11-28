import fitz
from docx import Document
from pathlib import Path

from app.domain.ports.file_loader import FileLoaderPort


class DocFileLoader(FileLoaderPort):
    async def read(self, path: str) -> str:
        return await _load_file(path)


async def _load_file(file_path: str) -> str:
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    text = ""

    if path.suffix.lower() == ".pdf":
        try:
            with fitz.open(path) as doc:
                for page in doc:
                    page_text = page.get_text()
                    if page_text:
                        text += page_text
        except Exception as exc:
            raise RuntimeError(f"PDF read failed: {exc}") from exc

    elif path.suffix.lower() == ".docx":
        try:
            doc = Document(path)
            text_parts: list[str] = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            text_parts.append(cell_text)

            text = "\n".join(text_parts)
        except Exception as exc:
            raise RuntimeError(f"DOCX read failed: {exc}") from exc
    else:
        raise ValueError("Unsupported file type")

    if not text.strip():
        raise ValueError("File contains no readable text (possibly scanned or image-based document)")

    return text.strip()
