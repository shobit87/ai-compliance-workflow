import fitz
from docx import Document
from pathlib import Path


async def load_file(file_path: str) -> str:
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    text = ""

    # -------- PDF ---------
    if path.suffix.lower() == ".pdf":
        try:
            with fitz.open(path) as doc:
                for page in doc:
                    page_text = page.get_text()
                    if page_text:
                        text += page_text
        except Exception as e:
            raise RuntimeError(f"PDF read failed: {e}")

    # -------- DOCX --------
    elif path.suffix.lower() == ".docx":
        try:
            doc = Document(path)
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            text_parts.append(cell_text)

            text = "\n".join(text_parts)
        except Exception as e:
            raise RuntimeError(f"DOCX read failed: {e}")

    else:
        raise ValueError("Unsupported file type")

    # FINAL GUARD
    if not text.strip():
        raise ValueError("File contains no readable text (possibly scanned or image-based document)")

    return text.strip()
